"""
File processing service for handling text and document files
"""
import io
import logging
from typing import BinaryIO
from docx import Document

logger = logging.getLogger(__name__)


class FileProcessor:
    """Service for processing uploaded files"""
    
    async def process_file(self, file_content: bytes, filename: str) -> str:
        """
        Process uploaded file and extract text content
        
        Args:
            file_content: Raw file bytes
            filename: Name of the uploaded file
            
        Returns:
            Extracted text content
        """
        try:
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'txt':
                return await self._process_text_file(file_content)
            elif file_extension == 'docx':
                return await self._process_docx_file(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            raise
    
    async def _process_text_file(self, file_content: bytes) -> str:
        """Process plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    logger.info(f"Successfully decoded text file with {encoding} encoding")
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            logger.warning("Used utf-8 with error replacement for text file")
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise ValueError("Could not process text file")
    
    async def _process_docx_file(self, file_content: bytes) -> str:
        """Process Microsoft Word document"""
        try:
            # Create a file-like object from bytes
            file_stream = io.BytesIO(file_content)
            
            # Load the document
            doc = Document(file_stream)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            if not text_parts:
                raise ValueError("No text content found in document")
            
            # Join all text parts
            full_text = '\n\n'.join(text_parts)
            logger.info(f"Extracted text from DOCX file: {len(full_text)} characters")
            
            return self._clean_text(full_text)
            
        except Exception as e:
            logger.error(f"Error processing DOCX file: {str(e)}")
            raise ValueError("Could not process Word document. Please ensure it's a valid .docx file")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = []
        for line in text.split('\n'):
            cleaned_line = ' '.join(line.split())
            if cleaned_line:
                lines.append(cleaned_line)
        
        # Join lines with proper spacing
        cleaned_text = '\n'.join(lines)
        
        # Remove excessive newlines
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
        
        return cleaned_text.strip()